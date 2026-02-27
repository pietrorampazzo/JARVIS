from docx import Document
import sys

def create_docx(text, output_path):
    doc = Document()
    doc.add_heading('JARVIS CORE ROADMAP', 0)
    for line in text.split('\n'):
        if line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('###'):
            doc.add_heading(line.replace('###', '').strip(), level=3)
        else:
            doc.add_paragraph(line)
    doc.save(output_path)
    print(f"Created docx at {output_path}")

if __name__ == "__main__":
    with open('c:/Users/pietr/OneDrive/.vscode/JARVIS/CORE_ROADMAP.md', 'r', encoding='utf-8') as f:
        content = f.read()
    create_docx(content, 'c:/Users/pietr/OneDrive/.vscode/JARVIS/CORE_ROADMAP.docx')
