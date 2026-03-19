import pandas as pd
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import logging
from datetime import datetime
import shutil
import locale

# ======================================================================
# CONFIGURAÇÕES
# ======================================================================

# Configuração de localização para formatação de moeda
try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil.1252')
    except locale.Error:
        print("Warning: Não foi possível configurar localização brasileira")

# Caminhos dos arquivos
BASE_DIR = r"C:\Users\pietr\OneDrive\.vscode\arte_"
CAMINHO_PLANILHA = os.path.join(BASE_DIR, "DOWNLOADS", "master_heavy_4.xlsx")
CAMINHO_TEMPLATE_MASTER = os.path.join(BASE_DIR,"DOWNLOADS", "PROPOSTAS_GERADAS", "TEMPLATES", "TEMPLATE_MASTER_ORCAMENTO.docx")
CAMINHO_SAIDA_DIR = os.path.join(BASE_DIR, "DOWNLOADS" ,"PROPOSTAS_GERADAS")

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(BASE_DIR, "LOGS", "automacao_orcamentos.log"), encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ======================================================================
# FUNÇÕES AUXILIARES
# ======================================================================

def formatar_moeda(valor):
    """Formatar valor para moeda brasileira"""
    try:
        if pd.isna(valor) or valor == 0:
            return "R$ 0,00"
        return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return "R$ 0,00"

def limpar_texto(texto):
    """Limpar e formatar texto para o documento"""
    if pd.isna(texto):
        return "x"
    texto_limpo = str(texto).strip()
    return texto_limpo if texto_limpo != '' else "x"

def criar_diretorio_se_nao_existir(caminho):
    """Criar diretório se não existir"""
    if not os.path.exists(caminho):
        os.makedirs(caminho)
        logger.info(f"Diretório criado: {caminho}")

def encontrar_tabela_itens(doc):
    """
    Encontra a tabela de itens no documento.
    Procura por tabela que contenha as colunas específicas do orçamento.
    """
    for i, table in enumerate(doc.tables):
        if table.rows and len(table.rows) > 0:
            # Verificar se a primeira linha contém os cabeçalhos esperados
            primeira_linha = [cell.text.strip().upper() for cell in table.rows[0].cells]
            
            # Verificar se contém palavras-chave dos cabeçalhos
            palavras_chave = ['ITEM', 'QUANT', 'MARCA', 'MODELO', 'DESCRIÇÃO', 'VALOR']
            if any(palavra in ' '.join(primeira_linha) for palavra in palavras_chave):
                logger.info(f"Tabela de itens encontrada: índice {i}")
                return i, table
    
    logger.warning("Tabela de itens não encontrada no template")
    return None, None

def atualizar_informacoes_pregao(doc, nome_pregao, data_atual):
    """
    Atualiza informações específicas do pregão no documento
    """
    try:
        # Percorrer todos os parágrafos procurando por placeholders
        for paragraph in doc.paragraphs:
            if '[PREGAO]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[PREGAO]', nome_pregao)
                logger.info(f"Placeholder [PREGAO] substituído por: {nome_pregao}")
            
            if '[DATA]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[DATA]', data_atual)
                logger.info(f"Placeholder [DATA] substituído por: {data_atual}")
        
        # Verificar também nas tabelas (caso haja placeholders em células)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '[PREGAO]' in cell.text:
                        cell.text = cell.text.replace('[PREGAO]', nome_pregao)
                    if '[DATA]' in cell.text:
                        cell.text = cell.text.replace('[DATA]', data_atual)
                        
    except Exception as e:
        logger.warning(f"Erro ao atualizar informações do pregão: {str(e)}")

def limpar_tabela_itens(table):
    """
    Remove todas as linhas da tabela exceto o cabeçalho
    """
    try:
        # Manter apenas a primeira linha (cabeçalho)
        linhas_para_remover = []
        for i in range(len(table.rows) - 1, 0, -1):  # Do final para o início
            linhas_para_remover.append(i)
        
        for i in linhas_para_remover:
            table._element.remove(table.rows[i]._element)
            
        logger.info("Tabela de itens limpa (mantido apenas cabeçalho)")
        
    except Exception as e:
        logger.error(f"Erro ao limpar tabela: {str(e)}")

def adicionar_linha_item(table, dados_item):
    """
    Adiciona uma linha com os dados do item na tabela
    """
    try:
        # Adicionar nova linha
        nova_linha = table.add_row()
        
        # Preencher células com os dados
        for i, dado in enumerate(dados_item):
            if i < len(nova_linha.cells):
                nova_linha.cells[i].text = str(dado)
                
                # Aplicar formatação específica baseada no conteúdo
                if i in [0, 1, 2, 6, 7]:  # ITEM, QUANT, UNIDADE, VALOR UNIT, VALOR TOTAL
                    nova_linha.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Destacar marca e modelo em negrito
                if i in [3, 4]:  # MARCA e MODELO
                    for paragraph in nova_linha.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            
    except Exception as e:
        logger.error(f"Erro ao adicionar linha: {str(e)}")

def adicionar_linha_total(table, total_geral):
    """
    Adiciona linha de total geral na tabela
    """
    try:
        linha_total = table.add_row()
        
        # Número de colunas da tabela
        num_cols = len(linha_total.cells)
        
        # Limpar todas as células
        for i in range(num_cols):
            linha_total.cells[i].text = ""
        
        # Preencher células específicas para o total
        if num_cols >= 8:
            linha_total.cells[-3].text = "TOTAL GERAL"  # Antepenúltima coluna
            linha_total.cells[-1].text = formatar_moeda(total_geral)  # Última coluna
            
            # Aplicar formatação de negrito e centralização
            for i in [-3, -1]:
                for paragraph in linha_total.cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        
    except Exception as e:
        logger.error(f"Erro ao adicionar linha de total: {str(e)}")

# ======================================================================
# FUNÇÃO PRINCIPAL DE GERAÇÃO DO ORÇAMENTO
# ======================================================================

def criar_orcamento_pregao_template(df_pregao, nome_arquivo):
    """
    Cria um orçamento individual baseado no template master
    """
    try:
        logger.info(f"Iniciando criação do orçamento para: {nome_arquivo}")
        
        # Verificar se template master existe
        if not os.path.exists(CAMINHO_TEMPLATE_MASTER):
            logger.error(f"❌ Template master não encontrado: {CAMINHO_TEMPLATE_MASTER}")
            return False
        
        # Filtrar apenas itens com match encontrado ou parcial
        df_validos = df_pregao[df_pregao['STATUS'].isin(['Match Encontrado', 'Match Parcial (Sugestão)'])].copy()
        
        if df_validos.empty:
            logger.warning(f"Nenhum item válido encontrado para o pregão: {nome_arquivo}")
            return False
        
        # Copiar template master para novo arquivo
        caminho_saida = os.path.join(CAMINHO_SAIDA_DIR, f"ORCAMENTO_{nome_arquivo}.docx")
        shutil.copy2(CAMINHO_TEMPLATE_MASTER, caminho_saida)
        
        # Abrir documento copiado
        doc = Document(caminho_saida)
        
        # Atualizar informações do pregão
        data_atual = datetime.now().strftime("%d de %B de %Y")
        atualizar_informacoes_pregao(doc, nome_arquivo, data_atual)
        
        # Encontrar tabela de itens
        idx_tabela, tabela_itens = encontrar_tabela_itens(doc)
        
        if tabela_itens is None:
            logger.error("❌ Tabela de itens não encontrada no template")
            return False
        
        # Limpar tabela (manter apenas cabeçalho)
        limpar_tabela_itens(tabela_itens)
        
        # Adicionar itens
        total_geral = 0
        for idx, (_, item) in enumerate(df_validos.iterrows(), 1):
            # Extrair e tratar dados da planilha
            item_numero = item.get('Nº', 'x') if pd.notna(item.get('Nº')) else 'x'
            qtde = int(float(item.get('QTDE', 0))) if pd.notna(item.get('QTDE')) and str(item.get('QTDE')).strip() != '' else 0
            unidade = limpar_texto(item.get('UNID_FORN', 'x')) if pd.notna(item.get('UNID_FORN')) and str(item.get('UNID_FORN')).strip() != '' else 'x'
            marca = limpar_texto(item.get('MARCA_SUGERIDA', 'x')) if pd.notna(item.get('MARCA_SUGERIDA')) and str(item.get('MARCA_SUGERIDA')).strip() != '' else 'x'
            modelo = limpar_texto(item.get('MODELO_SUGERIDO', 'x')) if pd.notna(item.get('MODELO_SUGERIDO')) and str(item.get('MODELO_SUGERIDO')).strip() != '' else 'x'
            descricao = limpar_texto(item.get('DESCRICAO', 'x')) if pd.notna(item.get('DESCRICAO')) and str(item.get('DESCRICAO')).strip() != '' else 'x'
            
            # Calcular valores
            preco_unit = float(item.get('PRECO_FINAL_VENDA', 0)) if pd.notna(item.get('PRECO_FINAL_VENDA')) and str(item.get('PRECO_FINAL_VENDA')).strip() != '' else 0
            valor_total = preco_unit * qtde
            total_geral += valor_total
            
            # Dados da linha - mapeamento correto
            dados_linha = [
                str(item_numero),  # ITEM (usa Nº da planilha)
                str(qtde),  # QUANT. (usa QTDE da planilha)
                unidade,  # UNIDADE (usa UNID_FORN da planilha)
                marca,  # Marca (usa MARCA_SUGERIDA da planilha)
                modelo,  # MODELO (usa MODELO_SUGERIDO da planilha)
                descricao,  # DESCRIÇÃO (usa DESCRICAO da planilha)
                formatar_moeda(preco_unit),  # VALOR UNIT. (usa PRECO_FINAL_VENDA da planilha)
                formatar_moeda(valor_total)  # VALOR TOTAL (PRECO_FINAL_VENDA * QTDE)
            ]
            
            # Adicionar linha na tabela
            adicionar_linha_item(tabela_itens, dados_linha)
        
        # Adicionar linha de total
        adicionar_linha_total(tabela_itens, total_geral)
        
        # Salvar documento
        doc.save(caminho_saida)
        
        logger.info(f"✅ Orçamento criado com sucesso: {caminho_saida}")
        logger.info(f"   - Itens processados: {len(df_validos)}")
        logger.info(f"   - Valor total: {formatar_moeda(total_geral)}")
        
        # Log detalhado dos primeiros 3 itens para debug
        logger.info("   - Primeiros itens processados:")
        for idx, (_, item) in enumerate(df_validos.head(3).iterrows(), 1):
            item_num = item.get('Nº', 'x')
            marca = item.get('MARCA_SUGERIDA', 'x')
            modelo = item.get('MODELO_SUGERIDO', 'x') 
            logger.info(f"     • Item {item_num}: {marca} - {modelo}")
        
        return True
        
    except Exception as e:
        logger.error(f"❌ Erro ao criar orçamento para {nome_arquivo}: {str(e)}")
        return False

def processar_todos_pregoes():
    """Função principal que processa todos os pregões"""
    try:
        logger.info("="*60)
        logger.info("🚀 INICIANDO SISTEMA DE AUTOMAÇÃO DE ORÇAMENTOS")
        logger.info("="*60)
        
        # Verificar se o arquivo da planilha existe
        if not os.path.exists(CAMINHO_PLANILHA):
            logger.error(f"❌ Arquivo não encontrado: {CAMINHO_PLANILHA}")
            return
        
        # Verificar se template master existe
        if not os.path.exists(CAMINHO_TEMPLATE_MASTER):
            logger.error(f"❌ Template master não encontrado: {CAMINHO_TEMPLATE_MASTER}")
            logger.info("💡 DICA: Crie o arquivo template em:")
            logger.info(f"   {CAMINHO_TEMPLATE_MASTER}")
            logger.info("   - Use seu orçamento atual como base")
            logger.info("   - Adicione [PREGAO] e [DATA] onde necessário")
            logger.info("   - Mantenha a estrutura da tabela de itens")
            return
        
        # Carregar planilha
        df = pd.read_excel(CAMINHO_PLANILHA)
        logger.info(f"📊 Planilha carregada: {len(df)} linhas")
        
        # Criar diretório de saída
        criar_diretorio_se_nao_existir(CAMINHO_SAIDA_DIR)
        
        # Obter lista única de pregões
        pregoes_unicos = df['ARQUIVO'].dropna().unique()
        logger.info(f"🎯 Encontrados {len(pregoes_unicos)} pregões únicos")
        
        # Estatísticas
        sucessos = 0
        erros = 0
        
        # Processar cada pregão
        for i, pregao in enumerate(pregoes_unicos, 1):
            nome_limpo = str(pregao).replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_')
            logger.info(f"🔄 [{i}/{len(pregoes_unicos)}] Processando: {nome_limpo}")
            
            # Filtrar dados do pregão
            df_pregao = df[df['ARQUIVO'] == pregao].copy()
            
            # Contar itens válidos
            itens_validos = len(df_pregao[df_pregao['STATUS'].isin(['Match Encontrado', 'Match Parcial (Sugestão)'])])
            logger.info(f"   - Total de itens: {len(df_pregao)}")
            logger.info(f"   - Itens válidos: {itens_validos}")
            
            # Criar orçamento
            if criar_orcamento_pregao_template(df_pregao, nome_limpo):
                sucessos += 1
                print(f"✅ [{i}/{len(pregoes_unicos)}] {nome_limpo}")
            else:
                erros += 1
                print(f"❌ [{i}/{len(pregoes_unicos)}] {nome_limpo}")
        
        # Relatório final
        logger.info("="*60)
        logger.info("📈 RELATÓRIO FINAL DE PROCESSAMENTO")
        logger.info("="*60)
        logger.info(f"✅ Orçamentos criados com sucesso: {sucessos}")
        logger.info(f"❌ Erros: {erros}")
        logger.info(f"📊 Total processado: {sucessos + erros}")
        logger.info(f"📁 Arquivos salvos em: {CAMINHO_SAIDA_DIR}")
        logger.info("="*60)
        
        print("\n" + "="*60)
        print("🎉 PROCESSAMENTO CONCLUÍDO!")
        print(f"✅ Sucessos: {sucessos}")
        print(f"❌ Erros: {erros}")
        print(f"📁 Diretório: {CAMINHO_SAIDA_DIR}")
        print("="*60)
        
        # Mostrar alguns arquivos criados como exemplo
        if sucessos > 0:
            arquivos = os.listdir(CAMINHO_SAIDA_DIR)
            arquivos_docx = [f for f in arquivos if f.endswith('.docx')][:5]
            print("\n📄 Exemplos de arquivos criados:")
            for arquivo in arquivos_docx:
                print(f"   • {arquivo}")
            if len(arquivos_docx) < sucessos:
                print(f"   • ... e mais {sucessos - len(arquivos_docx)} arquivos")
        
    except Exception as e:
        logger.error(f"❌ Erro crítico no processamento: {str(e)}")
        print(f"❌ Erro crítico: {str(e)}")

# ======================================================================
# FUNÇÃO PARA CRIAR TEMPLATE DE EXEMPLO
# ======================================================================

def criar_template_exemplo():
    """
    Cria um template de exemplo baseado no layout mostrado
    """
    try:
        # Criar diretório do template
        template_dir = os.path.dirname(CAMINHO_TEMPLATE_MASTER)
        criar_diretorio_se_nao_existir(template_dir)
        
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Cabeçalho com logos (simulado)
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run("🎵 ARTE COMERCIAL 🎵")
        run.bold = True
        run.font.size = Pt(16)
        
        # Informações do pregão
        doc.add_paragraph("")
        info_p = doc.add_paragraph(f"Pregão Eletrônico: [PREGAO]")
        info_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        data_p = doc.add_paragraph(f"Data: [DATA]")
        data_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Espaçamento
        doc.add_paragraph("")
        
        # Tabela de termos comerciais
        doc.add_paragraph("TERMOS COMERCIAIS", style='Heading 2')
        
        tabela_termos = doc.add_table(rows=2, cols=5)
        tabela_termos.style = 'Table Grid'
        
        headers_termos = ['Termos de Envio', 'Método de Envio', 'Prazo de Entrega', 'Cond. De Pagamento', 'Validade da Proposta']
        valores_termos = ['CIF', 'Transportadora', '30 D.', '30 D.', '60 D.']
        
        for i, (header, valor) in enumerate(zip(headers_termos, valores_termos)):
            tabela_termos.rows[0].cells[i].text = header
            tabela_termos.rows[1].cells[i].text = valor
            
            # Formatar cabeçalho
            for paragraph in tabela_termos.rows[0].cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
            
            # Centralizar valores
            tabela_termos.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espaçamento
        doc.add_paragraph("")
        
        # Tabela de itens (que será manipulada pelo script)
        doc.add_paragraph("ITENS DO ORÇAMENTO", style='Heading 2')
        
        tabela_itens = doc.add_table(rows=1, cols=8)
        tabela_itens.style = 'Table Grid'
        
        # Cabeçalhos da tabela de itens
        headers_itens = ['ITEM', 'QUANT.', 'UNIDADE', 'MARCA', 'MODELO', 'DESCRIÇÃO', 'VALOR UNIT.', 'VALOR TOTAL']
        
        for i, header in enumerate(headers_itens):
            cell = tabela_itens.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True
        
        # Adicionar observações no final
        doc.add_paragraph("")
        doc.add_paragraph("OBSERVAÇÕES:", style='Heading 3')
        obs_p = doc.add_paragraph()
        obs_p.add_run("• Nos preços acima estão inclusos todos os impostos, fretes, taxas, descarga e quaisquer outras que incidam direta ou indiretamente no fornecimento dos materiais desta licitação.\n")
        obs_p.add_run("• Os preços são válidos até a data da sua aceitação.\n")
        obs_p.add_run("• O prazo da validade da proposta será inferior a 60 (sessenta) dias a partir da data da sua apresentação.")
        
        # Salvar template
        doc.save(CAMINHO_TEMPLATE_MASTER)
        logger.info(f"✅ Template de exemplo criado: {CAMINHO_TEMPLATE_MASTER}")
        
        print("\n📝 TEMPLATE DE EXEMPLO CRIADO!")
        print(f"📁 Local: {CAMINHO_TEMPLATE_MASTER}")
        print("\n💡 PRÓXIMOS PASSOS:")
        print("1. Abra o template criado")
        print("2. Adicione seus logos e cabeçalhos reais")
        print("3. Ajuste as cores e formatação")
        print("4. Mantenha os placeholders [PREGAO] e [DATA]")
        print("5. Execute novamente o script para gerar os orçamentos")
        
        return True
        
    except Exception as e:
        logger.error(f"❌ Erro ao criar template: {str(e)}")
        return False

# ======================================================================
# EXECUÇÃO PRINCIPAL
# ======================================================================

if __name__ == "__main__":
    print("🚀 SISTEMA DE AUTOMAÇÃO DE ORÇAMENTOS - TEMPLATE BASED")
    print("="*65)
    
    # Verificar se existe template master
    if not os.path.exists(CAMINHO_TEMPLATE_MASTER):
        print("❓ Template master não encontrado!")
        print("\nOpções:")
        print("1. Criar template de exemplo")
        print("2. Sair e criar template manualmente")
        
        escolha = input("\nDigite sua escolha (1 ou 2): ").strip()
        
        if escolha == "1":
            criar_template_exemplo()
        else:
            print(f"\n💡 Crie seu template em: {CAMINHO_TEMPLATE_MASTER}")
            print("   - Use seu orçamento atual como base")
            print("   - Adicione placeholders [PREGAO] e [DATA]")
            print("   - Mantenha a estrutura da tabela de itens")
    else:
        # Processar todos os pregões
        processar_todos_pregoes()
    
    print("\n🏁 Execução finalizada!")